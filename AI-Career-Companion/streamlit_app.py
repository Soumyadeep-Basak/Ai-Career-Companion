import os
import streamlit as st
from app.modules.generate_doc import generate_resume, generate_cover_letter
# from app.modules.mock_interview import run_mock_interview
from app.modules.cheat_sheet import generate_cheat_sheet, extract_topics_from_text, generate_combined_cheat_sheet
from app.modules.career_map import generate_career_map

st.set_page_config(page_title="AI Career Companion", layout="wide", page_icon="💼")
st.title("💼 AI Career Companion")
st.set_page_config(page_title="AI Career Companion", layout="wide", page_icon="💼")
st.title("💼 AI Career Companion")

st.sidebar.markdown("## 📂 Navigation")
sidebar_choice = st.sidebar.radio(
st.sidebar.markdown("## 📂 Navigation")
sidebar_choice = st.sidebar.radio(
    "Choose a Module",
    ["📄 Generate Docs", "🧠 Mock Interview", "📚 Cheat Sheet", "🗸️ Career Map"]
    ["📄 Generate Docs", "🧠 Mock Interview", "📚 Cheat Sheet", "🗸️ Career Map"]
)

if sidebar_choice == "📄 Generate Docs":
    st.header("📄 Resume & Cover Letter Generator")
    uploaded_resume = st.file_uploader("📄 Upload your resume (PDF)", type=["pdf"])
if sidebar_choice == "📄 Generate Docs":
    st.header("📄 Resume & Cover Letter Generator")
    uploaded_resume = st.file_uploader("📄 Upload your resume (PDF)", type=["pdf"])
    resume_details = None
    jd = st.text_area("📋 Paste Job Description", key='jd', height=200)
    col1, col2 = st.columns(2)
    with col1:
        generate_resume_clicked = st.button("✨ Generate Resume", use_container_width=True)
    with col2:
        generate_cover_letter_clicked = st.button("📧 Generate Cover Letter", use_container_width=True)

    if uploaded_resume:
    jd = st.text_area("📋 Paste Job Description", key='jd', height=200)
    col1, col2 = st.columns(2)
    with col1:
        generate_resume_clicked = st.button("✨ Generate Resume", use_container_width=True)
    with col2:
        generate_cover_letter_clicked = st.button("📧 Generate Cover Letter", use_container_width=True)

    if uploaded_resume:
        from app.modules.generate_doc.resume_parser import extract_resume_details_from_pdf
        resume_details = extract_resume_details_from_pdf(uploaded_resume)
        st.text_area("📃 Extracted Resume Text", resume_details.get('raw_text', ''), height=200)
        st.text_area("📃 Extracted Resume Text", resume_details.get('raw_text', ''), height=200)

    if generate_resume_clicked:
        with st.spinner("⏳ Generating your resume..."):
            resume = generate_resume(jd, resume_details)
            st.markdown("""
<style>
.resume-box {
    background: #f9f9f9;
    border-radius: 8px;
    padding: 24px 32px;
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 1.05rem;
    line-height: 1.7;
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
    margin-bottom: 1.5em;
}
.resume-box ul { margin-left: 1.5em; }
.resume-box li { margin-bottom: 0.5em; }
.resume-box h2, .resume-box h3, .resume-box h4 { margin-top: 1.2em; }
</style>
<div class="resume-box">""", unsafe_allow_html=True)
            st.markdown(resume.replace('\n', '<br>'), unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)
            st.session_state['show_suggestions'] = True
            st.session_state['resume'] = resume

    if generate_cover_letter_clicked:
        improved_resume_text = st.session_state.get('improved_resume_text') or st.session_state.get('resume') or (resume_details.get('raw_text', '') if resume_details else '')
        import re
        name = ''
        for line in improved_resume_text.split('\n'):
            if len(re.findall(r'[A-Z][a-z]+', line)) >= 2 and len(line.split()) <= 5:
                name = line.strip()
                break
        cover_letter_prompt = f"Write a professional cover letter for the following job description using this resume.\n\nJob Description:\n{jd}\n\nResume:\n{improved_resume_text}\n\nMy name is {name}."
        from app.modules.generate_doc import genai
        model = genai.GenerativeModel('models/gemini-1.5-flash-latest')
        cl_response = model.generate_content(cover_letter_prompt)
        cl = cl_response.text if hasattr(cl_response, 'text') else str(cl_response)
        st.text_area("📨 Generated Cover Letter", cl, height=300)

        from io import BytesIO
        import docx
        doc = docx.Document()
        doc.add_heading('Cover Letter', 0)
        for para in cl.split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip(), style='Normal')
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(
            label="📅 Download Cover Letter (DOCX)",
            data=buffer,
            file_name="Cover_Letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    if st.session_state.get('show_suggestions', False) and not generate_cover_letter_clicked:
        st.subheader("🔧 Upgrade Suggestions")
        with st.expander("⚙️ Advanced Options"):
            upgrade_projects = st.checkbox("Suggest Projects", key='upgrade_projects')
            upgrade_skills = st.checkbox("Suggest Skills", key='upgrade_skills')
        with st.spinner("⏳ Generating your resume..."):
            resume = generate_resume(jd, resume_details)
            st.markdown("""
<style>
.resume-box {
    background: #f9f9f9;
    border-radius: 8px;
    padding: 24px 32px;
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 1.05rem;
    line-height: 1.7;
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
    margin-bottom: 1.5em;
}
.resume-box ul { margin-left: 1.5em; }
.resume-box li { margin-bottom: 0.5em; }
.resume-box h2, .resume-box h3, .resume-box h4 { margin-top: 1.2em; }
</style>
<div class="resume-box">""", unsafe_allow_html=True)
            st.markdown(resume.replace('\n', '<br>'), unsafe_allow_html=True)
            st.markdown("</div>", unsafe_allow_html=True)
            st.session_state['show_suggestions'] = True
            st.session_state['resume'] = resume

    if generate_cover_letter_clicked:
        improved_resume_text = st.session_state.get('improved_resume_text') or st.session_state.get('resume') or (resume_details.get('raw_text', '') if resume_details else '')
        import re
        name = ''
        for line in improved_resume_text.split('\n'):
            if len(re.findall(r'[A-Z][a-z]+', line)) >= 2 and len(line.split()) <= 5:
                name = line.strip()
                break
        cover_letter_prompt = f"Write a professional cover letter for the following job description using this resume.\n\nJob Description:\n{jd}\n\nResume:\n{improved_resume_text}\n\nMy name is {name}."
        from app.modules.generate_doc import genai
        model = genai.GenerativeModel('models/gemini-1.5-flash-latest')
        cl_response = model.generate_content(cover_letter_prompt)
        cl = cl_response.text if hasattr(cl_response, 'text') else str(cl_response)
        st.text_area("📨 Generated Cover Letter", cl, height=300)

        from io import BytesIO
        import docx
        doc = docx.Document()
        doc.add_heading('Cover Letter', 0)
        for para in cl.split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip(), style='Normal')
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button(
            label="📅 Download Cover Letter (DOCX)",
            data=buffer,
            file_name="Cover_Letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    if st.session_state.get('show_suggestions', False) and not generate_cover_letter_clicked:
        st.subheader("🔧 Upgrade Suggestions")
        with st.expander("⚙️ Advanced Options"):
            upgrade_projects = st.checkbox("Suggest Projects", key='upgrade_projects')
            upgrade_skills = st.checkbox("Suggest Skills", key='upgrade_skills')
        suggestions = None
        if (st.session_state.get('upgrade_projects') or st.session_state.get('upgrade_skills')) and resume_details:
            from app.modules.generate_doc import genai
            prompt = f"Given the following resume and job description, suggest improvements.\n\nResume:\n{resume_details.get('raw_text', '')}\n\nJob Description:\n{jd}"
            if st.session_state.get('upgrade_projects'):
                prompt += "\n\nSuggest 2-3 impactful projects to add or upgrade."
                prompt += "\n\nSuggest 2-3 impactful projects to add or upgrade."
            if st.session_state.get('upgrade_skills'):
                prompt += "\n\nSuggest 2-3 in-demand skills to learn or improve."
            prompt += "\n\nIf any important certifications are missing, recommend them as well."
                prompt += "\n\nSuggest 2-3 in-demand skills to learn or improve."
            prompt += "\n\nIf any important certifications are missing, recommend them as well."
            model = genai.GenerativeModel('models/gemini-1.5-flash-latest')
            response = model.generate_content(prompt)
            suggestions = response.text if hasattr(response, 'text') else str(response)
            st.text_area("💡 Suggestions", suggestions, height=200)

            if suggestions and st.button("🔄 Apply Suggestions to Resume"):
                full_prompt = f"Rewrite my resume for this job: {jd}\n\nCurrent resume:\n{resume_details.get('raw_text', '')}\n\nApply these suggestions: {suggestions}\n\nNote: Add suggested projects, skills, and certifications."
                improved_resume = model.generate_content(full_prompt)
                improved_resume_text = improved_resume.text if hasattr(improved_resume, 'text') else str(improved_resume)
                st.session_state['improved_resume_text'] = improved_resume_text
                st.markdown("""
<style>
.resume-box {
    background: #f9f9f9;
    border-radius: 8px;
    padding: 24px 32px;
    font-family: 'Segoe UI', Arial, sans-serif;
    font-size: 1.05rem;
    line-height: 1.7;
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
    margin-bottom: 1.5em;
}
.resume-box ul { margin-left: 1.5em; }
.resume-box li { margin-bottom: 0.5em; }
.resume-box h2, .resume-box h3, .resume-box h4 { margin-top: 1.2em; }
</style>
<div class="resume-box">""", unsafe_allow_html=True)
                st.markdown(improved_resume_text.replace('\n', '<br>'), unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)

                from io import BytesIO
                import re
                import docx

                clean_text = re.sub(r'[★*•●▪️-]+', '', improved_resume_text)
                clean_text = re.sub(r'\n+', '\n', clean_text)
                doc = docx.Document()
                doc.add_heading('Professional Resume', 0)
                for para in clean_text.split('\n'):
                    if para.strip():
                        doc.add_paragraph(para.strip(), style='Normal')
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="📅 Download DOCX Resume",
                    data=buffer,
                    file_name="ATS_Optimized_Resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

elif sidebar_choice == "🧠 Mock Interview":
    st.title("🎤 Voice-Based Mock Interview")
    st.markdown("### Step 1: Upload Resume and JD")
    # Resume Upload
    uploaded_resume = st.file_uploader("📄 Upload Resume (PDF)", type=["pdf"], key="mock_resume")
    jd_text = st.text_area("📌 Paste Job Description", key="mock_jd")
    resume_text = ""
    if uploaded_resume:
        from app.modules.generate_doc.resume_parser import extract_resume_details_from_pdf
        parsed = extract_resume_details_from_pdf(uploaded_resume)
        resume_text = parsed.get("raw_text", "")

    if st.button("🧠 Generate Questions from Resume + JD") and (jd_text or resume_text):
        from app.modules.voice_interview import get_questions_from_resume_and_jd
        question_list = get_questions_from_resume_and_jd(resume_text, jd_text)
        st.session_state["questions"] = question_list
        st.success("✅ Questions generated successfully!")

    if "questions" in st.session_state:
        selected = st.selectbox(
            "📋 Choose a question to answer:",
            [q["question"] for q in st.session_state["questions"]],
        )
        st.session_state["selected_question"] = selected

    if st.session_state.get("selected_question"):
        st.markdown("### 🎯 Selected Question")
        st.info(st.session_state["selected_question"])

        st.markdown("### 📁 Upload your recorded answer (WebM or WAV)")

        # Upload audio file
        audio_file = st.file_uploader("Upload audio file", type=["webm", "wav","mp3"])

        if audio_file is not None:
            from app.modules.voice_interview.record_audio import save_uploaded_audio
            ext = os.path.splitext(audio_file.name)[1]
            audio_path = save_uploaded_audio(audio_file, extension=ext)

            # Play back audio
            st.audio(audio_path, format=f"audio/{ext[1:]}")
            st.success(f"✅ Audio uploaded and saved to: {audio_path}")

            # Evaluate button
            if st.button("💡 Evaluate Answer"):
                from app.modules.voice_interview import generate_feedback_from_audio
                result = generate_feedback_from_audio(audio_path, st.session_state["selected_question"])

                st.subheader("📝 Transcribed Answer:")
                st.text_area("Transcript", result["transcript"], height=200)

                st.subheader("📊 Gemini Evaluation:")
                st.markdown(result["feedback"], unsafe_allow_html=True)


elif sidebar_choice == "📚 Cheat Sheet":
    st.header("📚 Algorithm Cheat Sheet Generator")
    st.markdown("### 🧠 Auto-generate from Resume + Job Description")
elif sidebar_choice == "📚 Cheat Sheet":
    st.header("📚 Algorithm Cheat Sheet Generator")
    st.markdown("### 🧠 Auto-generate from Resume + Job Description")
    uploaded_resume = st.file_uploader("📄 Upload your resume (PDF)", type=["pdf"], key="cs_resume")
    jd_text = st.text_area("📝 Paste Job Description", key="cs_jd")

    extracted_text = ""
    if uploaded_resume:
        from app.modules.generate_doc.resume_parser import extract_resume_details_from_pdf
        resume_details = extract_resume_details_from_pdf(uploaded_resume)
        extracted_text = resume_details.get('raw_text', '')

    if st.button("🚀 Generate Cheat Sheets"):
        if not extracted_text.strip() and not jd_text.strip():
            st.warning("⚠️ Please upload a resume and/or paste a job description.")
    if st.button("🚀 Generate Cheat Sheets"):
        if not extracted_text.strip() and not jd_text.strip():
            st.warning("⚠️ Please upload a resume and/or paste a job description.")
        else:
            combined_text = f"{extracted_text}\n\n{jd_text}"
            topics = extract_topics_from_text(extracted_text, jd_text)
            combined_text = f"{extracted_text}\n\n{jd_text}"
            topics = extract_topics_from_text(extracted_text, jd_text)
            st.success(f"🔍 Detected Topics: {', '.join(topics)}")
            with st.spinner("🛠️ Generating cheat sheets..."):
            with st.spinner("🛠️ Generating cheat sheets..."):
                for topic in topics:
                    sheet = generate_cheat_sheet(topic, jd_text)
                    with st.expander(f"📌 {topic}", expanded=False):
                    sheet = generate_cheat_sheet(topic, jd_text)
                    with st.expander(f"📌 {topic}", expanded=False):
                        st.markdown(sheet, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### ✍️ Generate by Custom Topic")
    st.markdown("### ✍️ Generate by Custom Topic")
    custom_topic = st.text_input("Enter a specific algorithm topic (e.g., Binary Search Trees)")
    if st.button("📄 Generate Custom Cheat Sheet"):
        sheet = generate_cheat_sheet(custom_topic, jd_text)
    if st.button("📄 Generate Custom Cheat Sheet"):
        sheet = generate_cheat_sheet(custom_topic, jd_text)
        with st.expander(f"📌 Cheat Sheet: {custom_topic}", expanded=True):
            st.markdown(sheet, unsafe_allow_html=True)

elif sidebar_choice == "🗸️ Career Map":
    st.header("🗸️ Career Path Explorer")
    profile = st.text_area("🧑‍🎓 Tell us about your interests & skills")
    if st.button("🌟 Generate Career Map"):
        with st.spinner("🧭 Mapping your career journey..."):
            roadmap = generate_career_map(profile)
            st.text_area("📌 Career Roadmap", roadmap, height=400)
elif sidebar_choice == "🗸️ Career Map":
    st.header("🗸️ Career Path Explorer")
    profile = st.text_area("🧑‍🎓 Tell us about your interests & skills")
    if st.button("🌟 Generate Career Map"):
        with st.spinner("🧭 Mapping your career journey..."):
            roadmap = generate_career_map(profile)
            st.text_area("📌 Career Roadmap", roadmap, height=400)

st.markdown("---")
st.markdown("<center><small>🚀 Built with ❤️ using Streamlit | AI Career Companion</small></center>", unsafe_allow_html=True)
st.markdown("---")
st.markdown("<center><small>🚀 Built with ❤️ using Streamlit | AI Career Companion</small></center>", unsafe_allow_html=True)